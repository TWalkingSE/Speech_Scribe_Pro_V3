#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
📤 Sistema de Exportação Avançado - Speech Scribe Pro V3
Suporte a múltiplos formatos de exportação
"""

import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from dataclasses import dataclass

from speech_scribe.utils.logger import logger


@dataclass
class ExportResult:
    """Resultado de uma exportação"""
    success: bool
    output_path: str
    format: str
    size_bytes: int
    error: Optional[str] = None


class BaseExporter(ABC):
    """Classe base para exportadores"""
    
    @property
    @abstractmethod
    def format_name(self) -> str:
        """Nome do formato"""
        pass
    
    @property
    @abstractmethod
    def file_extension(self) -> str:
        """Extensão do arquivo"""
        pass
    
    @abstractmethod
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        """
        Exporta dados para o formato.
        
        Args:
            data: Dados da transcrição
            output_path: Caminho de saída
            **options: Opções específicas do formato
        """
        pass
    
    def _ensure_extension(self, path: Path) -> Path:
        """Garante que o arquivo tem a extensão correta"""
        if path.suffix.lower() != self.file_extension.lower():
            path = path.with_suffix(self.file_extension)
        return path


class TXTExporter(BaseExporter):
    """Exportador para texto simples"""
    
    @property
    def format_name(self) -> str:
        return "Plain Text"
    
    @property
    def file_extension(self) -> str:
        return ".txt"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            output_path = self._ensure_extension(output_path)
            
            text = data.get('text', '')
            
            # Incluir timestamps se solicitado
            if options.get('include_timestamps', False) and 'segments' in data:
                lines = []
                for seg in data['segments']:
                    start = seg.get('start', 0)
                    text_seg = seg.get('text', '').strip()
                    lines.append(f"[{start:.2f}s] {text_seg}")
                text = '\n'.join(lines)
            
            output_path.write_text(text, encoding='utf-8')
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class JSONExporter(BaseExporter):
    """Exportador para JSON"""
    
    @property
    def format_name(self) -> str:
        return "JSON"
    
    @property
    def file_extension(self) -> str:
        return ".json"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            output_path = self._ensure_extension(output_path)
            
            # Adicionar metadados
            export_data = {
                'metadata': {
                    'exported_at': datetime.now().isoformat(),
                    'format_version': '1.0',
                    'generator': 'Speech Scribe Pro V3'
                },
                'transcription': data
            }
            
            indent = options.get('indent', 2)
            content = json.dumps(export_data, ensure_ascii=False, indent=indent, default=str)
            output_path.write_text(content, encoding='utf-8')
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class SRTExporter(BaseExporter):
    """Exportador para legendas SRT"""
    
    @property
    def format_name(self) -> str:
        return "SubRip (SRT)"
    
    @property
    def file_extension(self) -> str:
        return ".srt"
    
    def _format_timestamp(self, seconds: float) -> str:
        """Formata timestamp no padrão SRT: HH:MM:SS,mmm"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            output_path = self._ensure_extension(output_path)
            
            segments = data.get('segments', [])
            
            if not segments:
                # Criar segmento único se não houver segmentos
                segments = [{'start': 0, 'end': 0, 'text': data.get('text', '')}]
            
            lines = []
            for i, seg in enumerate(segments, 1):
                start = self._format_timestamp(seg.get('start', 0))
                end = self._format_timestamp(seg.get('end', seg.get('start', 0) + 3))
                text = seg.get('text', '').strip()
                
                lines.append(str(i))
                lines.append(f"{start} --> {end}")
                lines.append(text)
                lines.append("")
            
            output_path.write_text('\n'.join(lines), encoding='utf-8')
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class VTTExporter(BaseExporter):
    """Exportador para legendas WebVTT"""
    
    @property
    def format_name(self) -> str:
        return "WebVTT"
    
    @property
    def file_extension(self) -> str:
        return ".vtt"
    
    def _format_timestamp(self, seconds: float) -> str:
        """Formata timestamp no padrão VTT: HH:MM:SS.mmm"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            output_path = self._ensure_extension(output_path)
            
            segments = data.get('segments', [])
            
            lines = ["WEBVTT", ""]
            
            # Metadados
            if options.get('include_metadata', True):
                lines.append(f"NOTE Generated by Speech Scribe Pro V3")
                lines.append(f"NOTE Date: {datetime.now().isoformat()}")
                lines.append("")
            
            for i, seg in enumerate(segments, 1):
                start = self._format_timestamp(seg.get('start', 0))
                end = self._format_timestamp(seg.get('end', seg.get('start', 0) + 3))
                text = seg.get('text', '').strip()
                
                # Cue identifier (opcional)
                if options.get('include_cue_ids', True):
                    lines.append(f"cue-{i}")
                
                lines.append(f"{start} --> {end}")
                lines.append(text)
                lines.append("")
            
            output_path.write_text('\n'.join(lines), encoding='utf-8')
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class DOCXExporter(BaseExporter):
    """Exportador para Microsoft Word"""
    
    @property
    def format_name(self) -> str:
        return "Microsoft Word"
    
    @property
    def file_extension(self) -> str:
        return ".docx"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error="python-docx não instalado. Execute: pip install python-docx"
            )
        
        try:
            output_path = self._ensure_extension(output_path)
            
            doc = Document()
            
            # Título
            title = doc.add_heading('Transcrição', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadados
            if options.get('include_metadata', True):
                doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
                doc.add_paragraph(f"Idioma: {data.get('language', 'N/A')}")
                if 'model_used' in data:
                    doc.add_paragraph(f"Modelo: {data['model_used']}")
                doc.add_paragraph()
            
            # Conteúdo
            if options.get('include_timestamps', False) and 'segments' in data:
                doc.add_heading('Transcrição com Timestamps', level=1)
                for seg in data['segments']:
                    start = seg.get('start', 0)
                    text = seg.get('text', '').strip()
                    speaker = seg.get('speaker', '')
                    
                    if speaker:
                        p = doc.add_paragraph()
                        p.add_run(f"[{speaker}] ").bold = True
                        p.add_run(f"({start:.1f}s) ").italic = True
                        p.add_run(text)
                    else:
                        doc.add_paragraph(f"[{start:.1f}s] {text}")
            else:
                doc.add_heading('Texto', level=1)
                doc.add_paragraph(data.get('text', ''))
            
            # Estatísticas
            if options.get('include_stats', True):
                doc.add_heading('Estatísticas', level=1)
                text = data.get('text', '')
                words = text.split()
                doc.add_paragraph(f"Palavras: {len(words)}")
                doc.add_paragraph(f"Caracteres: {len(text)}")
                if 'processing_time' in data:
                    doc.add_paragraph(f"Tempo de processamento: {data['processing_time']:.1f}s")
            
            doc.save(str(output_path))
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class HTMLExporter(BaseExporter):
    """Exportador para HTML"""
    
    @property
    def format_name(self) -> str:
        return "HTML"
    
    @property
    def file_extension(self) -> str:
        return ".html"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            output_path = self._ensure_extension(output_path)
            
            template = options.get('template', 'default')
            
            html = self._generate_html(data, template, options)
            output_path.write_text(html, encoding='utf-8')
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )
    
    def _generate_html(self, data: Dict[str, Any], template: str, options: Dict) -> str:
        """Gera HTML da transcrição"""
        text = data.get('text', '')
        segments = data.get('segments', [])
        
        # CSS
        css = """
        body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; background: #1e1e1e; color: #d4d4d4; }
        h1 { color: #569cd6; border-bottom: 2px solid #569cd6; padding-bottom: 10px; }
        .metadata { background: #2d2d2d; padding: 15px; border-radius: 8px; margin-bottom: 20px; }
        .metadata span { margin-right: 20px; }
        .segment { margin: 10px 0; padding: 10px; background: #252526; border-radius: 4px; }
        .timestamp { color: #6a9955; font-family: monospace; }
        .speaker { color: #ce9178; font-weight: bold; }
        .stats { display: flex; gap: 20px; flex-wrap: wrap; }
        .stat-card { background: #2d2d2d; padding: 15px; border-radius: 8px; min-width: 150px; }
        .stat-value { font-size: 24px; font-weight: bold; color: #4ec9b0; }
        """
        
        # Metadados
        meta_html = ""
        if options.get('include_metadata', True):
            meta_html = f"""
            <div class="metadata">
                <span>📅 {datetime.now().strftime('%d/%m/%Y %H:%M')}</span>
                <span>🌐 {data.get('language', 'N/A')}</span>
                <span>🤖 {data.get('model_used', 'N/A')}</span>
            </div>
            """
        
        # Conteúdo
        content_html = ""
        if segments and options.get('include_timestamps', True):
            for seg in segments:
                speaker = seg.get('speaker', '')
                speaker_html = f'<span class="speaker">[{speaker}]</span> ' if speaker else ''
                content_html += f"""
                <div class="segment">
                    {speaker_html}
                    <span class="timestamp">[{seg.get('start', 0):.1f}s]</span>
                    {seg.get('text', '').strip()}
                </div>
                """
        else:
            content_html = f"<p>{text}</p>"
        
        # Estatísticas
        stats_html = ""
        if options.get('include_stats', True):
            words = text.split()
            stats_html = f"""
            <h2>📊 Estatísticas</h2>
            <div class="stats">
                <div class="stat-card">
                    <div class="stat-value">{len(words)}</div>
                    <div>Palavras</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">{len(text)}</div>
                    <div>Caracteres</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">{data.get('processing_time', 0):.1f}s</div>
                    <div>Processamento</div>
                </div>
            </div>
            """
        
        return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Transcrição - Speech Scribe Pro V3</title>
    <style>{css}</style>
</head>
<body>
    <h1>🎙️ Transcrição</h1>
    {meta_html}
    <h2>📝 Conteúdo</h2>
    {content_html}
    {stats_html}
    <footer style="margin-top: 40px; color: #666; text-align: center;">
        Gerado por Speech Scribe Pro V3
    </footer>
</body>
</html>"""


class CSVExporter(BaseExporter):
    """Exportador para CSV (segmentos)"""
    
    @property
    def format_name(self) -> str:
        return "CSV"
    
    @property
    def file_extension(self) -> str:
        return ".csv"
    
    def export(self, data: Dict[str, Any], output_path: Path, **options) -> ExportResult:
        try:
            import csv
            
            output_path = self._ensure_extension(output_path)
            
            segments = data.get('segments', [])
            
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f, delimiter=options.get('delimiter', ','))
                
                # Cabeçalho
                headers = ['index', 'start', 'end', 'text']
                if any('speaker' in seg for seg in segments):
                    headers.insert(3, 'speaker')
                writer.writerow(headers)
                
                # Dados
                for i, seg in enumerate(segments):
                    row = [
                        i + 1,
                        f"{seg.get('start', 0):.3f}",
                        f"{seg.get('end', 0):.3f}",
                    ]
                    if 'speaker' in headers:
                        row.append(seg.get('speaker', ''))
                    row.append(seg.get('text', '').strip())
                    writer.writerow(row)
            
            return ExportResult(
                success=True,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=output_path.stat().st_size
            )
        except Exception as e:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=self.format_name,
                size_bytes=0,
                error=str(e)
            )


class ExportManager:
    """Gerenciador de exportação"""
    
    def __init__(self):
        self._exporters: Dict[str, BaseExporter] = {}
        self._register_default_exporters()
    
    def _register_default_exporters(self):
        """Registra exportadores padrão"""
        self.register_exporter('txt', TXTExporter())
        self.register_exporter('json', JSONExporter())
        self.register_exporter('srt', SRTExporter())
        self.register_exporter('vtt', VTTExporter())
        self.register_exporter('html', HTMLExporter())
        self.register_exporter('csv', CSVExporter())
        self.register_exporter('docx', DOCXExporter())
    
    def register_exporter(self, format_key: str, exporter: BaseExporter):
        """Registra um exportador"""
        self._exporters[format_key.lower()] = exporter
        logger.debug(f"Exportador registrado: {format_key}")
    
    def get_exporter(self, format_key: str) -> Optional[BaseExporter]:
        """Obtém exportador por formato"""
        return self._exporters.get(format_key.lower())
    
    def get_available_formats(self) -> List[Dict[str, str]]:
        """Lista formatos disponíveis"""
        return [
            {
                'key': key,
                'name': exp.format_name,
                'extension': exp.file_extension
            }
            for key, exp in self._exporters.items()
        ]
    
    def export(self, data: Dict[str, Any], output_path: Path, 
               format_key: str, **options) -> ExportResult:
        """
        Exporta dados.
        
        Args:
            data: Dados da transcrição
            output_path: Caminho de saída
            format_key: Formato de exportação
            **options: Opções do formato
        """
        exporter = self.get_exporter(format_key)
        
        if exporter is None:
            return ExportResult(
                success=False,
                output_path=str(output_path),
                format=format_key,
                size_bytes=0,
                error=f"Formato não suportado: {format_key}"
            )
        
        result = exporter.export(data, output_path, **options)
        
        if result.success:
            logger.info(f"✅ Exportado: {result.output_path} ({result.format})")
        else:
            logger.error(f"❌ Erro na exportação: {result.error}")
        
        return result


# Instância singleton
_export_manager: Optional[ExportManager] = None


def get_export_manager() -> ExportManager:
    """Obtém gerenciador de exportação singleton"""
    global _export_manager
    if _export_manager is None:
        _export_manager = ExportManager()
    return _export_manager


__all__ = [
    'BaseExporter',
    'TXTExporter',
    'JSONExporter',
    'SRTExporter',
    'VTTExporter',
    'DOCXExporter',
    'HTMLExporter',
    'CSVExporter',
    'ExportResult',
    'ExportManager',
    'get_export_manager',
]
